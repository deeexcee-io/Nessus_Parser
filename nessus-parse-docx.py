import xml.etree.ElementTree as ET
from docx import Document

def create_word_report(nessus_file, output_file):
    """Parse the Nessus XML file and create a Word report with IP addresses and vulnerabilities."""
    tree = ET.parse(nessus_file)
    root = tree.getroot()

    document = Document()
    
    for report in root.findall("./Report"):
        for host in report.findall("./ReportHost"):
            host_ip = host.find("HostProperties/tag[@name='host-ip']").text

            vulnerabilities = []
            for item in host.findall("./ReportItem"):
                severity = item.find("risk_factor").text
                if severity in ["Critical", "High", "Medium", "Low"]:
                    plugin_name = item.get("pluginName")
                    plugin_output = item.find("plugin_output")
                    port = item.get("port")			

                    # Split the plugin name at the first ':'
                    plugin_name_parts = plugin_name.split(":", 1)
                    if len(plugin_name_parts) > 1:
                        # Take everything after the ':' character
                        plugin_name = plugin_name_parts[1].strip()
                    
                    info = f"[+] Plugin Name: {plugin_name}\n"
                    if plugin_output is not None and port is not None:
                        info += f"[+] Port Number: {port}\n[+] Output: {plugin_output.text}\n"
                    elif port is not None:
                        info += f"[+] Port Number: {port}\n"
                    elif plugin_output is not None:
                        info += f"[+] Output: {plugin_output.text}"
                    else:
                        info += "No Plugin Output or Port Information available."

                    vulnerabilities.append(info)

            if vulnerabilities:
                document.add_heading("IP Address: " + host_ip, level=1)
                for vulnerability in vulnerabilities:
                    document.add_paragraph(vulnerability)
                document.add_page_break()
    
    document.save(output_file)
    print(f"Word report saved to {output_file}")

if __name__ == '__main__':
    nessus_file = input("Enter the path to your Nessus file: ")
    output_file = input("Enter the output Word document file name (e.g., report.docx): ")
    create_word_report(nessus_file, output_file)
